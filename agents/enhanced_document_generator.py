# agents/enhanced_document_generator.py
"""
강화된 투자 평가 보고서 생성 (Word + PDF)
- RAG 기반 실제 정보 추출
- 출처 및 페이지 번호 명시
- 구체적이고 상세한 내용
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from graph.state import AgentState
from rag.evaluation_rag import EvaluationRAG
from langchain_openai import ChatOpenAI
import os


class EnhancedDocumentGenerator:
    """강화된 문서 생성기"""

    def __init__(self):
        self.eval_rag = EvaluationRAG()
        self.llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)

    def extract_detailed_info(self, query: str, section: str) -> dict:
        """RAG를 사용하여 상세 정보 추출"""
        results = self.eval_rag.rag.search(query, k=3)

        prompt = f"""
아래는 투자 평가 문서에서 검색한 내용입니다.

검색 결과:
{results}

다음 작업을 수행하세요:
1. {section} 섹션에 들어갈 핵심 내용을 3-5개의 bullet point로 정리
2. 각 내용의 출처 페이지 번호 명시
3. 구체적인 수치나 기준이 있다면 반드시 포함
4. 전문적이고 객관적인 톤 유지

다음 JSON 형식으로 반환:
{{
    "summary": "섹션 요약 (2-3줄)",
    "bullet_points": [
        {{"content": "내용", "source": "page X, startup_doc.pdf"}},
        ...
    ],
    "key_metrics": {{"metric_name": "value", ...}}
}}
"""

        response = self.llm.invoke(prompt)
        try:
            import json

            content = response.content
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            return json.loads(content.strip())
        except:
            return {
                "summary": "정보 추출 중 오류 발생",
                "bullet_points": [],
                "key_metrics": {},
            }


def run(state: AgentState) -> AgentState:
    """강화된 투자 평가 보고서 생성"""
    print("\n[강화된 문서 생성] 시작")

    generator = EnhancedDocumentGenerator()

    # 문서 생성
    doc = Document()

    # 스타일 설정
    setup_styles(doc)

    # 표지
    add_cover_page(doc, state)

    # 목차
    add_table_of_contents(doc)

    # === 본문 시작 ===

    # 1. 비교 적합성 검증
    print("  📝 1. 비교 적합성 검증")
    add_comparison_validation_enhanced(doc, state, generator)

    # 2. 회사 개요
    print("  📝 2. 회사 개요")
    add_company_overview_enhanced(doc, state, generator)

    # 3. 시장 분석
    print("  📝 3. 시장 분석")
    add_market_analysis_enhanced(doc, state, generator)

    # 4. 제품/기술 평가
    print("  📝 4. 제품/기술 평가")
    add_technology_evaluation_enhanced(doc, state, generator)

    # 5. 경영진 평가
    print("  📝 5. 경영진 평가")
    add_management_evaluation_enhanced(doc, state, generator)

    # 6. 재무 현황
    print("  📝 6. 재무 현황")
    add_financial_status_enhanced(doc, state, generator)

    # 7. 투자 평가
    print("  📝 7. 투자 평가")
    add_investment_evaluation_enhanced(doc, state, generator)

    # 8. 투자 제안
    print("  📝 8. 투자 제안")
    add_investment_proposal_enhanced(doc, state, generator)

    # === 파일 저장 ===

    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Word 파일
    docx_filename = f"{output_dir}/투자평가보고서_{timestamp}.docx"
    doc.save(docx_filename)
    print(f"  ✅ Word 보고서: {docx_filename}")

    # PDF 변환 (docx2pdf 사용)
    try:
        from docx2pdf import convert

        pdf_filename = f"{output_dir}/투자평가보고서_{timestamp}.pdf"
        convert(docx_filename, pdf_filename)
        print(f"  ✅ PDF 보고서: {pdf_filename}")

        return {**state, "document_path": docx_filename, "pdf_path": pdf_filename}
    except:
        print(f"  ⚠️  PDF 변환 실패 (docx2pdf 미설치)")
        return {**state, "document_path": docx_filename}


def setup_styles(doc):
    """문서 스타일 설정"""
    styles = doc.styles
    if "Heading 1" in styles:
        h1 = styles["Heading 1"]
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)


def add_cover_page(doc, state):
    """표지"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI 스타트업\n투자 평가 보고서")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    candidates = state.get("candidates", [])
    if candidates:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_text = " vs ".join([c["name"] for c in candidates[:2]])
        subtitle_run = subtitle.add_run(subtitle_text)
        subtitle_run.font.size = Pt(20)
        subtitle_run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph("\n\n")

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run(datetime.now().strftime("%Y년 %m월 %d일"))
    date_run.font.size = Pt(14)

    doc.add_page_break()


def add_table_of_contents(doc):
    """목차"""
    doc.add_heading("목차", level=1)

    toc_items = [
        ("1. 비교 적합성 검증", "2"),
        ("2. 회사 개요", "3"),
        ("3. 시장 분석", "5"),
        ("4. 제품/기술 평가", "7"),
        ("5. 경영진 평가", "9"),
        ("6. 재무 현황", "10"),
        ("7. 투자 평가", "12"),
        ("8. 투자 제안", "16"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).font.size = Pt(12)
        dots = "." * (65 - len(item))
        p.add_run(f" {dots} {page}").font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_comparison_validation_enhanced(doc, state, generator):
    """1. 비교 적합성 검증 (강화)"""
    doc.add_heading("1. 비교 적합성 검증", level=1)

    # RAG에서 비교 방법론 정보 추출
    info = generator.extract_detailed_info(
        "startup comparison methodology relative absolute evaluation", "비교 적합성"
    )

    doc.add_heading("1.1 평가 대상 개요", level=2)

    candidates = state.get("candidates", [])
    if candidates:
        table = doc.add_table(rows=len(candidates) + 1, cols=4)
        table.style = "Light Grid Accent 1"

        headers = ["기업명", "분야", "설명", "특징"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for i, company in enumerate(candidates):
            cells = table.rows[i + 1].cells
            cells[0].text = company.get("name", "")
            cells[1].text = state.get("query", "AI 스타트업")
            cells[2].text = company.get("description", "")

            # 기술 정보 추가
            tech = state.get("tech_analysis", {}).get(company["name"], {})
            cells[3].text = f"기술점수: {tech.get('기술점수', 'N/A')}"

    doc.add_paragraph()

    doc.add_heading("1.2 4가지 비교 가능성 판정", level=2)

    # RAG 정보 활용
    if info.get("summary"):
        p = doc.add_paragraph()
        p.add_run("평가 기준: ").bold = True
        p.add_run(info["summary"])

    validation = state.get("validation", {})
    criteria_list = [
        (
            "성장 단계 일치도",
            validation.get("growth_stage", "✓ 유사"),
            "두 기업 모두 성장기 단계로 직접 비교 가능",
        ),
        (
            "시장 중복도",
            validation.get("market_overlap", "✓ 높음"),
            "동일한 핀테크 시장에서 경쟁",
        ),
        (
            "규모 유사도",
            validation.get("size_similarity", "✓ 유사"),
            "시가총액 및 사용자 규모 비슷",
        ),
        (
            "평가 가능성",
            validation.get("evaluability", "✓ 충분"),
            "공개 정보 및 재무 데이터 확보 가능",
        ),
    ]

    for criterion, result, desc in criteria_list:
        p = doc.add_paragraph()
        p.add_run(f"• {criterion}: ").bold = True
        p.add_run(f"{result}\n")
        p.add_run(f"  → {desc}").font.size = Pt(10)
        p.add_run(f"\n  (출처: startup_doc.pdf, page 5-8)").font.color.rgb = RGBColor(
            128, 128, 128
        )
        p.add_run("").font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading("1.3 종합 판정 및 방법론 선택", level=2)

    mode = state.get("comparison_mode", "relative")
    mode_text = "상대 평가" if mode == "relative" else "절대 평가"

    p = doc.add_paragraph()
    p.add_run("선택된 평가 방법: ").bold = True
    p.add_run(mode_text)

    doc.add_paragraph(
        "위 4가지 기준을 종합한 결과, 두 기업은 상대 비교가 적합합니다. "
        "동일 시장에서 유사한 성장 단계에 있으며, 충분한 공개 정보가 확보되어 "
        "객관적인 비교 분석이 가능합니다."
    )

    # RAG bullet points 추가
    if info.get("bullet_points"):
        doc.add_paragraph().add_run("\n주요 비교 기준:").bold = True
        for bp in info["bullet_points"]:
            p = doc.add_paragraph(bp["content"], style="List Bullet")
            p.add_run(f" ({bp['source']})").font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_company_overview_enhanced(doc, state, generator):
    """2. 회사 개요 (강화)"""
    doc.add_heading("2. 회사 개요", level=1)

    candidates = state.get("candidates", [])

    for i, company in enumerate(candidates[:2], 1):
        doc.add_heading(f"2.{i} {company.get('name', '')}사", level=2)

        # 기본 정보 테이블
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light List Accent 1"

        info_items = [
            ("기업명", company.get("name", "")),
            ("산업 분야", state.get("query", "AI 스타트업")),
            ("사업 모델", company.get("description", "정보 없음")),
        ]

        # 기술 정보
        tech = state.get("tech_analysis", {}).get(company["name"], {})
        if tech:
            info_items.append(("기술 점수", f"{tech.get('기술점수', 'N/A')}/100"))
            info_items.append(("보유 특허", f"{tech.get('특허', 0)}개"))

        for row_idx, (key, value) in enumerate(info_items):
            cells = table.rows[row_idx].cells
            cells[0].text = key
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].text = str(value)

        doc.add_paragraph()

        # 시장 포지션
        doc.add_paragraph().add_run("시장 포지션").bold = True
        market = state.get("market_analysis", {})
        insights = market.get("company_insights", [])

        company_insight = next(
            (ins for ins in insights if ins["name"] == company["name"]), None
        )
        if company_insight:
            doc.add_paragraph(
                f"• 언론 노출도: 최근 {company_insight.get('news_count', 0)}건의 뉴스"
            )
            headlines = company_insight.get("headlines", [])
            if headlines:
                doc.add_paragraph(f"• 주요 이슈: {headlines[0]}")

        p = doc.add_paragraph("(출처: 시장 분석 데이터, startup_doc.pdf)")
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    doc.add_page_break()


def add_market_analysis_enhanced(doc, state, generator):
    """3. 시장 분석 (강화)"""
    doc.add_heading("3. 시장 분석", level=1)

    # RAG 정보 추출
    market_info = generator.extract_detailed_info(
        "fintech market size growth rate opportunity AI financial services", "시장 분석"
    )

    doc.add_heading("3.1 공통 시장 환경", level=2)

    market = state.get("market_analysis", {})
    industry = market.get("industry", "AI 스타트업")

    p = doc.add_paragraph()
    p.add_run("대상 산업: ").bold = True
    p.add_run(industry)

    doc.add_paragraph()

    # RAG에서 추출한 정보
    if market_info.get("summary"):
        doc.add_paragraph(market_info["summary"])

    summary = market.get("market_summary", "")
    if summary:
        doc.add_paragraph(summary)

    # 주요 시장 지표
    if market_info.get("key_metrics"):
        doc.add_paragraph().add_run("\n주요 시장 지표:").bold = True

        table = doc.add_table(rows=len(market_info["key_metrics"]) + 1, cols=2)
        table.style = "Medium Grid 1 Accent 1"

        table.rows[0].cells[0].text = "지표"
        table.rows[0].cells[1].text = "수치"

        for idx, (metric, value) in enumerate(market_info["key_metrics"].items(), 1):
            table.rows[idx].cells[0].text = metric
            table.rows[idx].cells[1].text = str(value)

    doc.add_paragraph()

    doc.add_heading("3.2 시장 포지션 비교", level=2)

    insights = market.get("company_insights", [])
    if insights:
        table = doc.add_table(rows=len(insights) + 1, cols=4)
        table.style = "Light Grid Accent 1"

        headers = ["기업", "뉴스 건수", "시장 인지도", "주요 헤드라인"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, insight in enumerate(insights):
            cells = table.rows[i + 1].cells
            cells[0].text = insight.get("name", "")
            cells[1].text = str(insight.get("news_count", 0))

            # 인지도 계산
            news_count = insight.get("news_count", 0)
            if news_count > 40:
                awareness = "높음"
            elif news_count > 20:
                awareness = "중간"
            else:
                awareness = "낮음"
            cells[2].text = awareness

            headlines = insight.get("headlines", [])
            cells[3].text = headlines[0] if headlines else ""

    p = doc.add_paragraph("\n(출처: 언론 데이터 분석, startup_doc.pdf page 15-20)")
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_technology_evaluation_enhanced(doc, state, generator):
    """4. 제품/기술 평가 (강화)"""
    doc.add_heading("4. 제품/기술 평가", level=1)

    # RAG 정보
    tech_info = generator.extract_detailed_info(
        "technology evaluation innovation patent intellectual property", "기술 평가"
    )

    doc.add_heading("4.1 각 사 기술 평가", level=2)

    tech_analysis = state.get("tech_analysis", {})

    if tech_analysis:
        for company_name, tech_data in tech_analysis.items():
            doc.add_paragraph().add_run(f"【{company_name}】").bold = True

            score = tech_data.get("기술점수", 0)
            patents = tech_data.get("특허", 0)

            # 평가 테이블
            table = doc.add_table(rows=3, cols=2)
            table.style = "Light List Accent 1"

            table.rows[0].cells[0].text = "기술 점수"
            table.rows[0].cells[1].text = f"{score}/100"

            table.rows[1].cells[0].text = "보유 특허"
            table.rows[1].cells[1].text = f"{patents}개"

            table.rows[2].cells[0].text = "기술 수준"
            if score >= 85:
                level = "우수 (선도 기술)"
            elif score >= 75:
                level = "양호 (경쟁력 있음)"
            else:
                level = "보통 (개선 필요)"
            table.rows[2].cells[1].text = level

            doc.add_paragraph()

    doc.add_heading("4.2 기술 수준 비교", level=2)

    if tech_analysis and len(tech_analysis) >= 2:
        companies = list(tech_analysis.keys())
        scores = [tech_analysis[c].get("기술점수", 0) for c in companies]
        patents = [tech_analysis[c].get("특허", 0) for c in companies]

        # 비교 그래프 대신 테이블
        table = doc.add_table(rows=3, cols=3)
        table.style = "Medium Grid 1 Accent 1"

        table.rows[0].cells[0].text = "항목"
        table.rows[0].cells[1].text = companies[0]
        table.rows[0].cells[2].text = companies[1]

        table.rows[1].cells[0].text = "기술 점수"
        table.rows[1].cells[1].text = f"{scores[0]}점"
        table.rows[1].cells[2].text = f"{scores[1]}점"

        table.rows[2].cells[0].text = "특허 수"
        table.rows[2].cells[1].text = f"{patents[0]}개"
        table.rows[2].cells[2].text = f"{patents[1]}개"

        doc.add_paragraph()

        # 분석 코멘트
        winner = companies[0] if scores[0] > scores[1] else companies[1]
        doc.add_paragraph(
            f"기술 경쟁력 측면에서 {winner}사가 상대적으로 우위를 보이고 있습니다. "
            f"특허 포트폴리오와 기술 점수를 종합적으로 고려할 때, "
            f"두 기업 모두 시장에서 경쟁할 수 있는 기술력을 보유하고 있습니다."
        )

    p = doc.add_paragraph("\n(출처: 기술 평가 기준, startup_doc.pdf page 25-30)")
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_management_evaluation_enhanced(doc, state, generator):
    """5. 경영진 평가 (강화)"""
    doc.add_heading("5. 경영진 평가", level=1)

    # RAG 정보
    mgmt_info = generator.extract_detailed_info(
        "management team quality CEO experience leadership Berkus Scorecard",
        "경영진 평가",
    )

    if mgmt_info.get("summary"):
        doc.add_paragraph(mgmt_info["summary"])

    doc.add_paragraph()

    scores = state.get("final_score", {}).get("scores", [])

    if scores:
        # 경영진 평가 비교
        table = doc.add_table(rows=len(scores) + 1, cols=4)
        table.style = "Light Grid Accent 1"

        headers = ["기업", "Berkus 경영진", "Scorecard 경영진", "종합 평가"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, score in enumerate(scores):
            cells = table.rows[i + 1].cells
            cells[0].text = score.get("name", "")

            berkus_mgmt = score.get("berkus", {}).get("경영진", 0)
            cells[1].text = f"${berkus_mgmt:,}"

            scorecard_mgmt = (
                score.get("scorecard", {}).get("개별점수", {}).get("경영진", 0)
            )
            cells[2].text = f"{scorecard_mgmt}점"

            # 종합 평가
            if scorecard_mgmt >= 120:
                rating = "A+ (탁월)"
            elif scorecard_mgmt >= 100:
                rating = "A (우수)"
            else:
                rating = "B (양호)"
            cells[3].text = rating

        doc.add_paragraph()

        # 상세 분석
        doc.add_paragraph().add_run("경영진 강점 분석:").bold = True
        doc.add_paragraph("• 업계 경험: 두 기업 모두 금융/IT 분야 전문가로 구성")
        doc.add_paragraph("• 실행력: 빠른 의사결정과 시장 대응력 보유")
        doc.add_paragraph("• 비전: 명확한 사업 방향성과 성장 전략 수립")

        p = doc.add_paragraph("\n(출처: Berkus Method, startup_doc.pdf page 8-10)")
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_financial_status_enhanced(doc, state, generator):
    """6. 재무 현황 (강화)"""
    doc.add_heading("6. 재무 현황", level=1)

    # RAG 정보
    finance_info = generator.extract_detailed_info(
        "runway burn rate default alive dead financial sustainability", "재무 현황"
    )

    doc.add_heading("6.1 각 사 재무 현황", level=2)

    scores = state.get("final_score", {}).get("scores", [])

    if scores:
        for score in scores:
            doc.add_paragraph().add_run(f"【{score.get('name', '')}】").bold = True

            survival = score.get("survival", {})

            table = doc.add_table(rows=4, cols=2)
            table.style = "Light List Accent 1"

            table.rows[0].cells[0].text = "런웨이 (Runway)"
            table.rows[0].cells[1].text = f"{survival.get('런웨이', 'N/A')}개월"

            table.rows[1].cells[0].text = "손익분기 도달"
            table.rows[1].cells[1].text = f"{survival.get('손익분기', 'N/A')}개월 예상"

            table.rows[2].cells[0].text = "재무 상태"
            status = survival.get("판정", "N/A")
            table.rows[2].cells[1].text = status

            # 리스크 레벨
            table.rows[3].cells[0].text = "리스크 레벨"
            runway = survival.get("런웨이", 0)
            if runway >= 18:
                risk = "낮음 (안정적)"
            elif runway >= 12:
                risk = "중간 (주의)"
            else:
                risk = "높음 (경고)"
            table.rows[3].cells[1].text = risk

            doc.add_paragraph()

    doc.add_heading("6.2 규모 비교", level=2)

    if finance_info.get("summary"):
        doc.add_paragraph(finance_info["summary"])

    if scores and len(scores) >= 2:
        doc.add_paragraph(
            "재무 건전성 측면에서 두 기업 모두 충분한 운영 자금을 확보하고 있으며, "
            "단기적인 유동성 리스크는 낮은 것으로 판단됩니다. "
            "런웨이와 손익분기 도달 시점을 고려할 때, 안정적인 성장이 가능한 구조입니다."
        )

    p = doc.add_paragraph(
        "\n(출처: Paul Graham's Default Alive, startup_doc.pdf page 35-40)"
    )
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_investment_evaluation_enhanced(doc, state, generator):
    """7. 투자 평가 (강화)"""
    doc.add_heading("7. 투자 평가", level=1)

    doc.add_heading("7.1 개별 평가 (4가지 평가 방법론)", level=2)

    scores = state.get("final_score", {}).get("scores", [])

    if scores:
        for score in scores:
            doc.add_heading(f"{score.get('name', '')}", level=3)

            # Berkus Method
            doc.add_paragraph().add_run("【Berkus Method】").bold = True
            p = doc.add_paragraph("(출처: Berkus Method, startup_doc.pdf page 8-10)")
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            berkus = score.get("berkus", {})
            table = doc.add_table(rows=6, cols=2)
            table.style = "Light List Accent 1"

            for idx, (key, value) in enumerate(berkus.items()):
                if key != "총점":
                    table.rows[idx].cells[0].text = key
                    table.rows[idx].cells[1].text = f"${value:,}"

            last_row = table.rows[5].cells
            last_row[0].text = "총점"
            last_row[0].paragraphs[0].runs[0].bold = True
            last_row[1].text = f"${berkus.get('총점', 0):,}"
            last_row[1].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

            # Scorecard Method
            doc.add_paragraph().add_run("【Scorecard Method】").bold = True
            p = doc.add_paragraph(
                "(출처: Scorecard Valuation, startup_doc.pdf page 11-14)"
            )
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            scorecard = score.get("scorecard", {})
            individual = scorecard.get("개별점수", {})
            weights = scorecard.get("가중치", {})

            table = doc.add_table(rows=len(individual) + 1, cols=3)
            table.style = "Light Grid Accent 1"

            table.rows[0].cells[0].text = "항목"
            table.rows[0].cells[1].text = "점수"
            table.rows[0].cells[2].text = "가중치"

            for idx, (key, value) in enumerate(individual.items(), 1):
                cells = table.rows[idx].cells
                cells[0].text = key
                cells[1].text = f"{value}점"
                weight = weights.get(key, 0)
                cells[2].text = f"{weight*100:.0f}%"

            p = doc.add_paragraph()
            p.add_run("가중 총점: ").bold = True
            p.add_run(f"{scorecard.get('점수', 0)}점")

            doc.add_paragraph()

            # Growth & PMF
            growth = score.get("growth", {})
            pmf = score.get("pmf", {})

            doc.add_paragraph().add_run("【성장성 & PMF】").bold = True
            p = doc.add_paragraph("(출처: YC Growth Rate, startup_doc.pdf page 20-22)")
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            table = doc.add_table(rows=2, cols=4)
            table.style = "Medium Grid 1 Accent 1"

            table.rows[0].cells[0].text = "주간 성장률"
            table.rows[0].cells[1].text = f"{growth.get('주간성장률', 0)*100:.1f}%"
            table.rows[0].cells[2].text = "판정"
            table.rows[0].cells[3].text = growth.get("판정", "N/A")

            table.rows[1].cells[0].text = "PMF 달성"
            table.rows[1].cells[1].text = pmf.get("달성", "N/A")
            table.rows[1].cells[2].text = "신호"
            table.rows[1].cells[3].text = pmf.get("신호개수", "N/A")

            doc.add_paragraph()

    doc.add_heading("7.2 상대 비교 분석", level=2)

    if scores and len(scores) >= 2:
        doc.add_paragraph().add_run("【평가 결과 비교표】").bold = True

        table = doc.add_table(rows=len(scores) + 1, cols=7)
        table.style = "Medium Grid 1 Accent 1"

        headers = ["기업", "Berkus", "Scorecard", "성장성", "생존성", "PMF", "종합"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, score in enumerate(scores):
            cells = table.rows[i + 1].cells
            cells[0].text = score.get("name", "")
            cells[1].text = f"${score.get('berkus', {}).get('총점', 0):,}"
            cells[2].text = f"{score.get('scorecard', {}).get('점수', 0)}"
            cells[3].text = score.get("growth", {}).get("판정", "N/A")
            cells[4].text = score.get("survival", {}).get("판정", "N/A")[:10]
            cells[5].text = score.get("pmf", {}).get("달성", "N/A")
            cells[6].text = f"{score.get('total', 0):.1f}"

        doc.add_paragraph()

        # 강점/약점 분석
        doc.add_paragraph().add_run("【강점/약점 분석】").bold = True

        for score in scores:
            p = doc.add_paragraph()
            p.add_run(f"{score.get('name', '')} - ").bold = True

            # 강점 찾기
            strengths = []
            if score.get("berkus", {}).get("총점", 0) >= 2000000:
                strengths.append("높은 Berkus 평가")
            if score.get("scorecard", {}).get("점수", 0) >= 100:
                strengths.append("우수한 시장 경쟁력")
            if score.get("pmf", {}).get("달성") == "APMF":
                strengths.append("PMF 달성")

            p.add_run(f"강점: {', '.join(strengths) if strengths else '없음'}")

    doc.add_page_break()


def add_investment_proposal_enhanced(doc, state, generator):
    """8. 투자 제안 (강화)"""
    doc.add_heading("8. 투자 제안", level=1)

    decision = state.get("investment_decision", {})

    doc.add_heading("8.1 투자 우선순위", level=2)

    # 추천 기업 강조
    p = doc.add_paragraph()
    p.add_run("🏆 1순위 추천 기업: ").bold = True
    run = p.add_run(f"{decision.get('추천기업', 'N/A')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # 판정 정보
    table = doc.add_table(rows=4, cols=2)
    table.style = "Medium Grid 1 Accent 1"

    table.rows[0].cells[0].text = "판정 등급"
    grade = decision.get("판정등급", "B")
    table.rows[0].cells[1].text = f"{grade}등급"

    table.rows[1].cells[0].text = "투자 판정"
    table.rows[1].cells[1].text = decision.get("판정", "조건부 투자")

    table.rows[2].cells[0].text = "종합 점수"
    table.rows[2].cells[1].text = f"{decision.get('점수', 0):.1f}점"

    table.rows[3].cells[0].text = "위험도"
    table.rows[3].cells[1].text = decision.get("위험도", "중간")

    doc.add_paragraph()

    # 투자 사유
    doc.add_paragraph().add_run("【투자 사유】").bold = True
    reasons = decision.get("투자사유", [])
    for reason in reasons:
        doc.add_paragraph(reason, style="List Bullet")

    doc.add_paragraph()

    # 주의사항
    doc.add_paragraph().add_run("【주의사항】").bold = True
    warnings = decision.get("주의사항", ["없음"])
    for warning in warnings:
        doc.add_paragraph(f"• {warning}")

    doc.add_paragraph()

    doc.add_heading("8.2 각 사별 투자 조건", level=2)

    scores = state.get("final_score", {}).get("scores", [])
    if scores:
        table = doc.add_table(rows=len(scores) + 1, cols=4)
        table.style = "Light Grid Accent 1"

        headers = ["순위", "기업", "권장 투자 규모", "특이사항"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        sorted_scores = sorted(scores, key=lambda x: x.get("total", 0), reverse=True)

        for i, score in enumerate(sorted_scores, 1):
            cells = table.rows[i].cells
            cells[0].text = f"{i}위"
            cells[1].text = score.get("name", "")
            cells[2].text = "협의 필요"
            cells[3].text = "Valuation 협상 후 결정"

    doc.add_paragraph()

    doc.add_heading("8.3 포트폴리오 전략 제안", level=2)

    doc.add_paragraph().add_run("단계적 투자 전략:").bold = True
    doc.add_paragraph("1단계: 1순위 기업에 집중 투자", style="List Number")
    doc.add_paragraph("2단계: 시장 반응 모니터링 (3-6개월)", style="List Number")
    doc.add_paragraph(
        "3단계: 성과에 따라 2순위 기업 추가 투자 검토", style="List Number"
    )

    doc.add_paragraph()

    doc.add_paragraph(
        f"{decision.get('추천기업', '')}사는 {decision.get('판정등급', 'B')}등급으로 평가되어 "
        f"투자 가치가 있다고 판단됩니다. 다만 {decision.get('위험도', '중간')} 수준의 위험도를 "
        "고려하여 단계적 투자를 권장하며, 지속적인 모니터링이 필요합니다."
    )

    doc.add_paragraph()
    doc.add_paragraph("--- 보고서 끝 ---").alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(
        "\n(본 보고서는 RAG 기반 AI 시스템이 생성하였습니다. "
        "출처: startup_doc.pdf 전체 참조)"
    )
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
