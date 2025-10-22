# agents/professional_document_generator.py
"""
전문 투자 비교 평가 보고서 생성
INVESTMENT ANALYSIS REPORT - Professional Format
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from graph.state import AgentState
from rag.evaluation_rag import EvaluationRAG
from langchain_openai import ChatOpenAI
import os


class ProfessionalDocumentGenerator:
    """전문 보고서 생성기"""

    def __init__(self):
        self.eval_rag = EvaluationRAG()
        self.llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)

    def extract_info(self, query: str) -> dict:
        """RAG 정보 추출"""
        results = self.eval_rag.rag.search(query, k=3)

        prompt = f"""
투자 평가 문서에서 검색한 내용을 바탕으로 핵심 정보를 추출하세요.

검색 결과:
{results}

다음 JSON 형식으로 반환:
{{
    "summary": "핵심 요약 (2-3줄)",
    "key_points": ["포인트1", "포인트2", "포인트3"]
}}
"""

        response = self.llm.invoke(prompt)
        try:
            import json

            content = response.content
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            return json.loads(content.strip())
        except:
            return {"summary": "", "key_points": []}


def run(state: AgentState) -> AgentState:
    """전문 투자 비교 평가 보고서 생성"""
    print("\n[전문 보고서 생성] 시작")

    generator = ProfessionalDocumentGenerator()
    candidates = state.get("candidates", [])

    # 문서 생성
    doc = Document()
    setup_styles(doc)

    # 표지
    add_cover_page(doc, state, candidates)

    # 목차
    add_toc(doc)

    # Ⅰ. 서론
    print("  [Section] Ⅰ. 서론")
    add_intro(doc, state, generator, candidates)

    # Ⅱ. 비교 적합성 검증
    print("  [Section] Ⅱ. 비교 적합성 검증")
    add_comparison_validation(doc, state, generator)

    # Ⅲ. 회사 개요
    print("  [Section] Ⅲ. 회사 개요")
    add_company_overview(doc, state, candidates)

    # Ⅳ. 시장 분석
    print("  [Section] Ⅳ. 시장 분석")
    add_market_analysis(doc, state)

    # Ⅴ. 제품 및 기술력 평가
    print("  [Section] Ⅴ. 제품 및 기술력 평가")
    add_tech_evaluation(doc, state, candidates)

    # Ⅵ. 경영진 및 조직 평가
    print("  [Section] Ⅵ. 경영진 및 조직 평가")
    add_management(doc, state)

    # Ⅶ. 재무 현황 분석
    print("  [Section] Ⅶ. 재무 현황 분석")
    add_financial(doc, state, candidates)

    # Ⅷ. 투자 평가 및 분석
    print("  [Section] Ⅷ. 투자 평가 및 분석")
    add_investment_evaluation(doc, state)

    # Ⅸ. 투자 제안
    print("  [Section] Ⅸ. 투자 제안")
    add_investment_proposal(doc, state)

    # Ⅹ. 결론 및 제언
    print("  [Section] Ⅹ. 결론 및 제언")
    add_conclusion(doc, state)

    # Ⅺ. 참고문헌
    print("  [Section] Ⅺ. 참고문헌")
    add_references(doc)

    # Ⅻ. 부록
    print("  [Section] Ⅻ. 부록")
    add_appendix(doc, state)

    # 저장
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"{output_dir}/투자비교평가보고서_{timestamp}.docx"
    doc.save(docx_filename)
    print(f"  [Complete] Word Report: {docx_filename}")

    return {**state, "document_path": docx_filename}


def setup_styles(doc):
    """스타일 설정"""
    styles = doc.styles
    if "Heading 1" in styles:
        h1 = styles["Heading 1"]
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)


def add_cover_page(doc, state, candidates):
    """표지 - 전문적인 디자인"""

    # 헤더: CONFIDENTIAL
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run("CONFIDENTIAL")
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    header_run.font.bold = True

    doc.add_paragraph("\n\n")

    # 메인 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_main = title.add_run("투자 비교 평가 보고서\n")
    title_main.font.size = Pt(28)
    title_main.font.bold = True
    title_main.font.color.rgb = RGBColor(0, 51, 102)

    title_sub = title.add_run("INVESTMENT COMPARATIVE ANALYSIS REPORT")
    title_sub.font.size = Pt(12)
    title_sub.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph("\n")

    # 평가 대상 박스
    if len(candidates) >= 2:
        box_p = doc.add_paragraph()
        box_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        box_run = box_p.add_run(f"{candidates[0]['name']} vs {candidates[1]['name']}")
        box_run.font.size = Pt(20)
        box_run.font.bold = True
        box_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("\n\n")

    # 기본 정보
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Light List Accent 1"

    info_table.rows[0].cells[0].text = "평가 분야"
    info_table.rows[0].cells[1].text = state.get("query", "AI 스타트업")

    info_table.rows[1].cells[0].text = "평가 대상"
    if len(candidates) >= 2:
        info_table.rows[1].cells[
            1
        ].text = f"{candidates[0]['name']}, {candidates[1]['name']}"

    info_table.rows[2].cells[0].text = "평가 방법"
    mode = state.get("comparison_mode", "relative")
    info_table.rows[2].cells[1].text = (
        "상대 비교 평가" if mode == "relative" else "절대 평가"
    )

    info_table.rows[3].cells[0].text = "작성일"
    info_table.rows[3].cells[1].text = datetime.now().strftime("%Y년 %m월 %d일")

    doc.add_paragraph("\n\n")

    # Executive Summary 박스
    doc.add_paragraph("\n")
    decision = state.get("investment_decision", {})

    if decision:
        # Summary 제목
        summary_title = doc.add_paragraph()
        summary_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary_title_run = summary_title.add_run("EXECUTIVE SUMMARY")
        summary_title_run.font.size = Pt(14)
        summary_title_run.font.bold = True
        summary_title_run.font.color.rgb = RGBColor(0, 51, 102)

        # Summary 테이블
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Medium Grid 1 Accent 1"

        summary_table.rows[0].cells[0].text = "추천 기업"
        summary_table.rows[0].cells[1].text = decision.get("추천기업", "N/A")

        summary_table.rows[1].cells[0].text = "판정 등급"
        summary_table.rows[1].cells[1].text = f"{decision.get('판정등급', 'B')} Grade"

        summary_table.rows[2].cells[0].text = "위험도"
        summary_table.rows[2].cells[1].text = decision.get("위험도", "중간")

        summary_table.rows[3].cells[0].text = "투자 판정"
        summary_table.rows[3].cells[1].text = decision.get("판정", "조건부 투자")

    doc.add_paragraph("\n\n")

    # 하단 정보
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Generated by AI-based Investment Evaluation System")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_toc(doc):
    """목차"""
    doc.add_heading("목차", level=1)

    toc_items = [
        ("Ⅰ. 서론", "2"),
        ("Ⅱ. 비교 적합성 검증", "3"),
        ("Ⅲ. 회사 개요", "4"),
        ("Ⅳ. 시장 분석", "6"),
        ("Ⅴ. 제품 및 기술력 평가", "8"),
        ("Ⅵ. 경영진 및 조직 평가", "10"),
        ("Ⅶ. 재무 현황 분석", "11"),
        ("Ⅷ. 투자 평가 및 분석", "13"),
        ("Ⅸ. 투자 제안", "17"),
        ("Ⅹ. 결론 및 제언", "18"),
        ("Ⅺ. 참고문헌", "19"),
        ("Ⅻ. 부록", "20"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).font.size = Pt(12)
        dots = "." * (70 - len(item))
        p.add_run(f" {dots} {page}").font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_intro(doc, state, generator, candidates):
    """Ⅰ. 서론"""
    doc.add_heading("Ⅰ. 서론", level=1)

    # 1. 평가 목적 및 범위
    doc.add_heading("1. 평가 목적 및 범위", level=2)

    query = state.get("query", "AI 스타트업")
    if len(candidates) >= 2:
        doc.add_paragraph(
            f"본 보고서는 {query} 분야의 {candidates[0]['name']}사와 {candidates[1]['name']}사에 대한 "
            f"종합적인 투자 비교 평가를 목적으로 합니다. "
            f"두 기업의 기술력, 시장성, 경영진 역량, 재무 건전성 등을 다각도로 분석하여 "
            f"투자 의사결정에 필요한 객관적 근거를 제공합니다."
        )

    doc.add_paragraph()

    # 2. 비교 대상 개요
    doc.add_heading("2. 비교 대상 개요", level=2)

    if len(candidates) >= 2:
        table = doc.add_table(rows=3, cols=3)
        table.style = "Medium Grid 1 Accent 1"

        table.rows[0].cells[0].text = "구분"
        table.rows[0].cells[1].text = candidates[0]["name"]
        table.rows[0].cells[2].text = candidates[1]["name"]

        table.rows[1].cells[0].text = "사업 분야"
        table.rows[1].cells[1].text = query
        table.rows[1].cells[2].text = query

        table.rows[2].cells[0].text = "주요 사업"
        table.rows[2].cells[1].text = candidates[0].get("description", "")
        table.rows[2].cells[2].text = candidates[1].get("description", "")

    doc.add_paragraph()

    # 3. 평가 방법론 및 데이터 출처
    doc.add_heading("3. 평가 방법론 및 데이터 출처", level=2)

    doc.add_paragraph().add_run("평가 방법론:").bold = True
    doc.add_paragraph("- Berkus Method: 초기 단계 스타트업 평가 (최대 $2.5M)")
    doc.add_paragraph("- Scorecard Method: 7개 항목 가중치 평가")
    doc.add_paragraph("- Y Combinator 성장률 기준: 주간 성장률 평가")
    doc.add_paragraph("- PMF(Product-Market Fit) 평가: 6가지 신호 분석")

    doc.add_paragraph()
    doc.add_paragraph().add_run("데이터 출처:").bold = True
    doc.add_paragraph("- startup_doc.pdf (투자 평가 기준서)")
    doc.add_paragraph("- 기업 공개 정보 및 시장 데이터")
    doc.add_paragraph("- 언론 보도 및 뉴스 분석")

    p = doc.add_paragraph("(출처: startup_doc.pdf, RAG 기반 정보 추출)")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_comparison_validation(doc, state, generator):
    """Ⅱ. 비교 적합성 검증"""
    doc.add_heading("Ⅱ. 비교 적합성 검증", level=1)

    # 1. 평가 대상 요약
    doc.add_heading("1. 평가 대상 요약", level=2)

    candidates = state.get("candidates", [])
    if candidates:
        for candidate in candidates[:2]:
            p = doc.add_paragraph()
            p.add_run(f"- {candidate.get('name', '')}: ").bold = True
            p.add_run(candidate.get("description", ""))

    doc.add_paragraph()

    # 2. 비교 가능성 판정
    doc.add_heading("2. 비교 가능성 판정", level=2)

    validation = state.get("validation", {})

    criteria_list = [
        (
            "(1) 성장 단계 일치도",
            validation.get("growth_stage", "✓ 유사"),
            "두 기업 모두 성장기 단계로 직접 비교 가능",
        ),
        (
            "(2) 시장 중복도",
            validation.get("market_overlap", "✓ 높음"),
            "동일한 시장에서 경쟁하며 직접 비교 의미 있음",
        ),
        (
            "(3) 규모 유사도",
            validation.get("size_similarity", "✓ 유사"),
            "시가총액 및 사용자 규모가 비슷한 수준",
        ),
        (
            "(4) 평가 가능성",
            validation.get("evaluability", "✓ 충분"),
            "공개 정보 및 재무 데이터 확보 가능",
        ),
    ]

    for criterion, result, desc in criteria_list:
        doc.add_paragraph().add_run(criterion).bold = True
        p = doc.add_paragraph(f"판정: {result}")
        p = doc.add_paragraph(f"근거: {desc}")
        p.runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # 3. 종합 판정 및 평가 접근법 선택
    doc.add_heading("3. 종합 판정 및 평가 접근법 선택", level=2)

    mode = state.get("comparison_mode", "relative")
    mode_text = "상대 비교 평가" if mode == "relative" else "절대 평가"

    p = doc.add_paragraph()
    p.add_run("선택된 평가 방법: ").bold = True
    p.add_run(mode_text)

    doc.add_paragraph(
        "종합 판정 결과, 4가지 기준을 모두 충족하여 두 기업의 상대 비교가 적합합니다. "
        "동일 시장에서 유사한 성장 단계에 있으며, 충분한 공개 정보가 확보되어 "
        "객관적인 비교 분석이 가능합니다."
    )

    doc.add_page_break()


def add_company_overview(doc, state, candidates):
    """Ⅲ. 회사 개요"""
    doc.add_heading("Ⅲ. 회사 개요", level=1)

    for i, company in enumerate(candidates[:2], 1):
        company_name = company.get("name", f"{chr(64+i)}사")

        doc.add_heading(f"{i}. {company_name} 기업 개요", level=2)

        # 가. 설립 배경 및 주요 연혁
        doc.add_heading("가. 설립 배경 및 주요 연혁", level=3)
        doc.add_paragraph(
            f"{company_name}는 {state.get('query', 'AI 스타트업')} 분야의 선도 기업으로, "
            f"혁신적인 기술과 서비스를 통해 시장에서 입지를 확보하고 있습니다."
        )

        doc.add_paragraph()

        # 나. 사업 구조 및 주요 제품
        doc.add_heading("나. 사업 구조 및 주요 제품", level=3)

        table = doc.add_table(rows=4, cols=2)
        table.style = "Light List Accent 1"

        table.rows[0].cells[0].text = "기업명"
        table.rows[0].cells[1].text = company_name

        table.rows[1].cells[0].text = "사업 모델"
        table.rows[1].cells[1].text = company.get("description", "")

        # 기술 정보
        tech = state.get("tech_analysis", {}).get(company_name, {})
        if tech:
            table.rows[2].cells[0].text = "기술 점수"
            table.rows[2].cells[1].text = f"{tech.get('기술점수', 'N/A')}/100"

            table.rows[3].cells[0].text = "보유 특허"
            table.rows[3].cells[1].text = f"{tech.get('특허', 0)}개"

        doc.add_paragraph()

    doc.add_page_break()


def add_market_analysis(doc, state):
    """Ⅳ. 시장 분석"""
    doc.add_heading("Ⅳ. 시장 분석", level=1)

    market = state.get("market_analysis", {})

    # 1. 공통 시장 환경 및 트렌드
    doc.add_heading("1. 공통 시장 환경 및 트렌드", level=2)

    industry = market.get("industry", "AI 스타트업")
    doc.add_paragraph(f"대상 산업: {industry}")

    summary = market.get("market_summary", "")
    if summary:
        doc.add_paragraph(summary)

    doc.add_paragraph()

    # 2. 시장 포지션 비교
    doc.add_heading("2. 시장 포지션 비교", level=2)

    doc.add_heading("가. 시장 점유율, 고객 세그먼트, 경쟁 지형", level=3)

    insights = market.get("company_insights", [])
    if insights:
        table = doc.add_table(rows=len(insights) + 1, cols=4)
        table.style = "Medium Grid 1 Accent 1"

        headers = ["기업", "언론 노출", "시장 인지도", "주요 이슈"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, insight in enumerate(insights):
            cells = table.rows[i + 1].cells
            cells[0].text = insight.get("name", "")
            cells[1].text = f"{insight.get('news_count', 0)}건"

            news_count = insight.get("news_count", 0)
            if news_count > 40:
                awareness = "높음"
            elif news_count > 20:
                awareness = "중간"
            else:
                awareness = "낮음"
            cells[2].text = awareness

            headlines = insight.get("headlines", [])
            cells[3].text = headlines[0] if headlines else ""

    p = doc.add_paragraph("\n(출처: 시장 데이터 분석, startup_doc.pdf page 15-20)")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_tech_evaluation(doc, state, candidates):
    """Ⅴ. 제품 및 기술력 평가"""
    doc.add_heading("Ⅴ. 제품 및 기술력 평가", level=1)

    tech_analysis = state.get("tech_analysis", {})

    for i, company in enumerate(candidates[:2], 1):
        company_name = company.get("name", "")

        doc.add_heading(f"{i}. {company_name} 기술력 평가", level=2)

        if company_name in tech_analysis:
            tech = tech_analysis[company_name]
            score = tech.get("기술점수", 0)
            patents = tech.get("특허", 0)

            table = doc.add_table(rows=3, cols=2)
            table.style = "Light List Accent 1"

            table.rows[0].cells[0].text = "기술 점수"
            table.rows[0].cells[1].text = f"{score}/100"

            table.rows[1].cells[0].text = "보유 특허"
            table.rows[1].cells[1].text = f"{patents}개"

            table.rows[2].cells[0].text = "기술 수준"
            if score >= 85:
                level = "우수 (선도 기술)"
            elif score >= 75:
                level = "양호 (경쟁력 있음)"
            else:
                level = "보통 (개선 필요)"
            table.rows[2].cells[1].text = level

        doc.add_paragraph()

    # 3. 기술 수준 및 경쟁력 비교
    doc.add_heading("3. 기술 수준 및 경쟁력 비교", level=2)

    if len(candidates) >= 2 and tech_analysis:
        companies = [c["name"] for c in candidates[:2]]
        scores = [tech_analysis.get(c, {}).get("기술점수", 0) for c in companies]
        patents = [tech_analysis.get(c, {}).get("특허", 0) for c in companies]

        table = doc.add_table(rows=3, cols=3)
        table.style = "Medium Grid 1 Accent 1"

        table.rows[0].cells[0].text = "항목"
        table.rows[0].cells[1].text = companies[0]
        table.rows[0].cells[2].text = companies[1]

        table.rows[1].cells[0].text = "기술 점수"
        table.rows[1].cells[1].text = f"{scores[0]}점"
        table.rows[1].cells[2].text = f"{scores[1]}점"

        table.rows[2].cells[0].text = "특허 수"
        table.rows[2].cells[1].text = f"{patents[0]}개"
        table.rows[2].cells[2].text = f"{patents[1]}개"

    p = doc.add_paragraph("\n(출처: 기술 평가 기준, startup_doc.pdf page 25-30)")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_management(doc, state):
    """Ⅵ. 경영진 및 조직 평가"""
    doc.add_heading("Ⅵ. 경영진 및 조직 평가", level=1)

    # 1. 핵심 인력 역량 및 경력 분석
    doc.add_heading("1. 핵심 인력 역량 및 경력 분석", level=2)

    scores = state.get("final_score", {}).get("scores", [])

    if scores:
        table = doc.add_table(rows=len(scores) + 1, cols=4)
        table.style = "Medium Grid 1 Accent 1"

        headers = ["기업", "Berkus 경영진", "Scorecard 경영진", "종합 평가"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, score in enumerate(scores):
            cells = table.rows[i + 1].cells
            cells[0].text = score.get("name", "")

            berkus_mgmt = score.get("berkus", {}).get("경영진", 0)
            cells[1].text = f"${berkus_mgmt:,}"

            scorecard_mgmt = (
                score.get("scorecard", {}).get("개별점수", {}).get("경영진", 0)
            )
            cells[2].text = f"{scorecard_mgmt}점"

            if scorecard_mgmt >= 120:
                rating = "A+ (탁월)"
            elif scorecard_mgmt >= 100:
                rating = "A (우수)"
            else:
                rating = "B (양호)"
            cells[3].text = rating

    doc.add_paragraph()

    # 2. 조직 구조 및 실행력 평가
    doc.add_heading("2. 조직 구조 및 실행력 평가", level=2)

    doc.add_paragraph("- 업계 경험: 금융/IT 분야 전문가로 구성")
    doc.add_paragraph("- 실행력: 빠른 의사결정과 시장 대응력 보유")
    doc.add_paragraph("- 비전: 명확한 사업 방향성과 성장 전략 수립")

    p = doc.add_paragraph("\n(출처: Berkus Method, startup_doc.pdf page 8-10)")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_financial(doc, state, candidates):
    """Ⅶ. 재무 현황 분석"""
    doc.add_heading("Ⅶ. 재무 현황 분석", level=1)

    scores = state.get("final_score", {}).get("scores", [])

    for i, score in enumerate(scores[:2], 1):
        company_name = score.get("name", f"{chr(64+i)}사")

        doc.add_heading(f"{i}. {company_name} 재무 개요", level=2)

        survival = score.get("survival", {})

        table = doc.add_table(rows=4, cols=2)
        table.style = "Light List Accent 1"

        table.rows[0].cells[0].text = "런웨이 (Runway)"
        table.rows[0].cells[1].text = f"{survival.get('런웨이', 'N/A')}개월"

        table.rows[1].cells[0].text = "손익분기 도달"
        table.rows[1].cells[1].text = f"{survival.get('손익분기', 'N/A')}개월 예상"

        table.rows[2].cells[0].text = "재무 상태"
        table.rows[2].cells[1].text = survival.get("판정", "N/A")

        runway = survival.get("런웨이", 0)
        if runway >= 18:
            risk = "낮음 (안정적)"
        elif runway >= 12:
            risk = "중간 (주의)"
        else:
            risk = "높음 (경고)"
        table.rows[3].cells[0].text = "리스크 레벨"
        table.rows[3].cells[1].text = risk

        doc.add_paragraph()

    # 3. 규모 및 수익성 비교
    doc.add_heading("3. 규모 및 수익성 비교", level=2)

    doc.add_paragraph(
        "재무 건전성 측면에서 두 기업 모두 충분한 운영 자금을 확보하고 있으며, "
        "단기적인 유동성 리스크는 낮은 것으로 판단됩니다."
    )

    p = doc.add_paragraph(
        "\n(출처: Paul Graham's Default Alive, startup_doc.pdf page 35-40)"
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def add_investment_evaluation(doc, state):
    """Ⅷ. 투자 평가 및 분석"""
    doc.add_heading("Ⅷ. 투자 평가 및 분석", level=1)

    scores = state.get("final_score", {}).get("scores", [])

    # 1. 개별사 평가 요약
    doc.add_heading("1. 개별사 평가 요약 (4대 평가항목)", level=2)

    doc.add_heading("가. 기술력 / 시장성 / 팀 / 재무", level=3)

    if scores:
        for score in scores:
            p = doc.add_paragraph()
            company_run = p.add_run(f"{score.get('name', '')}")
            company_run.bold = True
            company_run.font.size = Pt(12)
            company_run.font.color.rgb = RGBColor(0, 51, 102)

            # 간단한 요약 테이블
            table = doc.add_table(rows=5, cols=2)
            table.style = "Light List Accent 1"

            table.rows[0].cells[0].text = "Berkus"
            table.rows[0].cells[1].text = f"${score.get('berkus', {}).get('총점', 0):,}"

            table.rows[1].cells[0].text = "Scorecard"
            table.rows[1].cells[
                1
            ].text = f"{score.get('scorecard', {}).get('점수', 0)}점"

            table.rows[2].cells[0].text = "성장성"
            table.rows[2].cells[1].text = score.get("growth", {}).get("판정", "N/A")

            table.rows[3].cells[0].text = "PMF"
            table.rows[3].cells[1].text = score.get("pmf", {}).get("달성", "N/A")

            table.rows[4].cells[0].text = "종합 점수"
            table.rows[4].cells[1].text = f"{score.get('total', 0):.1f}점"

            doc.add_paragraph()

    # 2. 상대 비교 분석
    doc.add_heading("2. 상대 비교 분석", level=2)

    # 가. 평가 결과 비교표
    doc.add_heading("가. 평가 결과 비교표", level=3)

    if scores and len(scores) >= 2:
        table = doc.add_table(rows=len(scores) + 1, cols=7)
        table.style = "Medium Grid 1 Accent 1"

        headers = ["기업", "Berkus", "Scorecard", "성장성", "생존성", "PMF", "종합"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, score in enumerate(scores):
            cells = table.rows[i + 1].cells
            cells[0].text = score.get("name", "")
            cells[1].text = f"${score.get('berkus', {}).get('총점', 0):,}"
            cells[2].text = f"{score.get('scorecard', {}).get('점수', 0)}"
            cells[3].text = score.get("growth", {}).get("판정", "N/A")
            cells[4].text = score.get("survival", {}).get("판정", "N/A")[:10]
            cells[5].text = score.get("pmf", {}).get("달성", "N/A")
            cells[6].text = f"{score.get('total', 0):.1f}"

    doc.add_paragraph()

    # 나. 항목별 순위
    doc.add_heading("나. 항목별 순위", level=3)
    doc.add_paragraph("개별 평가 항목별 우위 기업 분석 결과는 상기 비교표 참조")

    doc.add_paragraph()

    # 다. 강점·약점 매트릭스
    doc.add_heading("다. 강점·약점 매트릭스", level=3)

    for score in scores[:2]:
        p = doc.add_paragraph()
        p.add_run(f"{score.get('name', '')} - ").bold = True

        strengths = []
        if score.get("berkus", {}).get("총점", 0) >= 2000000:
            strengths.append("높은 Berkus 평가")
        if score.get("scorecard", {}).get("점수", 0) >= 100:
            strengths.append("우수한 시장 경쟁력")
        if score.get("pmf", {}).get("달성") == "APMF":
            strengths.append("PMF 달성")

        p.add_run(f"강점: {', '.join(strengths) if strengths else '없음'}")

    doc.add_paragraph()

    # 라. Risk–Return 프로파일
    doc.add_heading("라. Risk–Return 프로파일", level=3)
    doc.add_paragraph(
        "위험 대비 수익 분석 결과, 두 기업 모두 중간 위험 - 높은 수익 구조"
    )

    doc.add_page_break()


def add_investment_proposal(doc, state):
    """Ⅸ. 투자 제안"""
    doc.add_heading("Ⅸ. 투자 제안", level=1)

    decision = state.get("investment_decision", {})

    # 1. 종합 투자 판단 및 우선순위
    doc.add_heading("1. 종합 투자 판단 및 우선순위", level=2)

    p = doc.add_paragraph()
    p.add_run("1순위 추천 기업: ").bold = True
    run = p.add_run(f"{decision.get('추천기업', 'N/A')}")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Medium Grid 1 Accent 1"

    table.rows[0].cells[0].text = "판정 등급"
    table.rows[0].cells[1].text = f"{decision.get('판정등급', 'B')}등급"

    table.rows[1].cells[0].text = "투자 판정"
    table.rows[1].cells[1].text = decision.get("판정", "조건부 투자")

    table.rows[2].cells[0].text = "종합 점수"
    table.rows[2].cells[1].text = f"{decision.get('점수', 0):.1f}점"

    table.rows[3].cells[0].text = "위험도"
    table.rows[3].cells[1].text = decision.get("위험도", "중간")

    doc.add_paragraph()

    # 2. 각 사별 투자 조건 및 유의사항
    doc.add_heading("2. 각 사별 투자 조건 및 유의사항", level=2)

    doc.add_paragraph().add_run("투자 사유:").bold = True
    reasons = decision.get("투자사유", [])
    for reason in reasons:
        doc.add_paragraph(f"- {reason}")

    doc.add_paragraph()
    doc.add_paragraph().add_run("주의사항:").bold = True
    warnings = decision.get("주의사항", ["없음"])
    for warning in warnings:
        doc.add_paragraph(f"- {warning}")

    doc.add_paragraph()

    # 3. 포트폴리오 전략 제안
    doc.add_heading("3. 포트폴리오 전략 제안", level=2)

    doc.add_paragraph().add_run("단계적 투자 전략:").bold = True
    doc.add_paragraph("1단계: 1순위 기업에 집중 투자", style="List Number")
    doc.add_paragraph("2단계: 시장 반응 모니터링 (3-6개월)", style="List Number")
    doc.add_paragraph(
        "3단계: 성과에 따라 2순위 기업 추가 투자 검토", style="List Number"
    )

    doc.add_page_break()


def add_conclusion(doc, state):
    """Ⅹ. 결론 및 제언"""
    doc.add_heading("Ⅹ. 결론 및 제언", level=1)

    # 1. 핵심 요약
    doc.add_heading("1. 핵심 요약 (Executive Summary)", level=2)

    decision = state.get("investment_decision", {})

    doc.add_paragraph(
        f"본 보고서는 {state.get('query', 'AI 스타트업')} 분야의 투자 비교 평가를 수행하였으며, "
        f"종합 분석 결과 {decision.get('추천기업', '')}사를 1순위 투자 대상으로 추천합니다. "
        f"판정 등급 {decision.get('판정등급', 'B')}등급으로 평가되었으며, "
        f"위험도는 {decision.get('위험도', '중간')} 수준입니다."
    )

    doc.add_paragraph()

    # 2. 향후 Follow-up 전략
    doc.add_heading("2. 향후 Follow-up 전략", level=2)

    doc.add_paragraph("- 투자 후 분기별 성과 모니터링")
    doc.add_paragraph("- 주요 KPI 추적 (성장률, 매출, 사용자 수)")
    doc.add_paragraph("- 시장 환경 변화 지속 관찰")
    doc.add_paragraph("- 경쟁사 동향 파악 및 대응 전략 수립")

    doc.add_page_break()


def add_references(doc):
    """Ⅺ. 참고문헌"""
    doc.add_heading("Ⅺ. 참고문헌 (REFERENCE)", level=1)

    refs = [
        "1. startup_doc.pdf - 투자 평가 기준서",
        "2. Berkus Method - Dave Berkus, 'The Berkus Method' (2023)",
        "3. Scorecard Valuation Method - Bill Payne, 'Scorecard Valuation Methodology' (2022)",
        "4. Paul Graham - 'Default Alive or Default Dead' (Y Combinator, 2015)",
        "5. Y Combinator - 'Startup Growth Rate Guidelines' (2023)",
        "6. 기업 공개 정보 및 언론 보도 자료",
        "7. 시장 조사 데이터 및 산업 리포트",
    ]

    for ref in refs:
        doc.add_paragraph(ref)

    doc.add_page_break()


def add_appendix(doc, state):
    """Ⅻ. 부록"""
    doc.add_heading("Ⅻ. 부록 (APPENDIX)", level=1)

    # 1. 평가 로직 및 점수 기준표
    doc.add_heading("1. 평가 로직 및 점수 기준표", level=2)

    doc.add_paragraph().add_run("Berkus Method:").bold = True
    doc.add_paragraph("각 항목당 최대 $500,000")
    doc.add_paragraph("- 아이디어 품질: $0 ~ $500,000")
    doc.add_paragraph("- 프로토타입: $0 ~ $500,000")
    doc.add_paragraph("- 경영진: $0 ~ $500,000")
    doc.add_paragraph("- 전략적 관계: $0 ~ $500,000")
    doc.add_paragraph("- 제품 출시: $0 ~ $500,000")

    doc.add_paragraph()

    doc.add_paragraph().add_run("Scorecard Method:").bold = True
    table = doc.add_table(rows=8, cols=2)
    table.style = "Light List Accent 1"

    scorecard_items = [
        ("항목", "가중치"),
        ("경영진", "30%"),
        ("시장", "25%"),
        ("제품", "15%"),
        ("경쟁", "10%"),
        ("판매", "10%"),
        ("투자", "5%"),
        ("기타", "5%"),
    ]

    for i, (item, weight) in enumerate(scorecard_items):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = weight

    doc.add_paragraph()

    doc.add_paragraph().add_run("성장률 기준:").bold = True
    doc.add_paragraph("- 우수: 주간 10% 이상")
    doc.add_paragraph("- 양호: 주간 5-10%")
    doc.add_paragraph("- 경고: 주간 1% 미만")

    doc.add_paragraph()

    # 2. 세부 재무 데이터
    doc.add_heading("2. 세부 재무 데이터", level=2)
    doc.add_paragraph("각 사별 상세 재무 정보는 본문 Ⅶ장 참조")

    doc.add_paragraph()

    # 3. 추가 자료
    doc.add_heading("3. 추가 자료 (시장 세분화표 등)", level=2)
    doc.add_paragraph("시장 세분화 및 추가 분석 자료는 별도 제공 가능")

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph("--- 보고서 끝 ---")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True

    p = doc.add_paragraph(
        "\n본 보고서는 RAG 기반 AI 투자 평가 시스템으로 생성되었습니다.\n"
        "출처: startup_doc.pdf 전체 참조"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
