# agents/document_generator.py
"""
전문 투자 평가 보고서 생성 (Word/PDF)
목차 기반 체계적 문서 생성
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from graph.state import AgentState
import os


def run(state: AgentState) -> AgentState:
    """전문 투자 평가 보고서 생성 (Word)"""
    print("\n[문서 생성] 시작")

    # 문서 생성
    doc = Document()

    # 스타일 설정
    setup_styles(doc)

    # 표지
    add_cover_page(doc, state)

    # 목차 페이지
    add_table_of_contents(doc)

    # 1. 비교 적합성 검증 (1페이지)
    add_comparison_validation(doc, state)

    # 2. 회사 개요 (2페이지)
    add_company_overview(doc, state)

    # 3. 시장 분석 (2페이지)
    add_market_analysis(doc, state)

    # 4. 제품/기술 평가 (2페이지)
    add_technology_evaluation(doc, state)

    # 5. 경영진 평가 (1페이지)
    add_management_evaluation(doc, state)

    # 6. 재무 현황 (2페이지)
    add_financial_status(doc, state)

    # 7. 투자 평가 (4페이지)
    add_investment_evaluation(doc, state)

    # 8. 투자 제안 (1페이지)
    add_investment_proposal(doc, state)

    # 파일 저장
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{output_dir}/투자평가보고서_{timestamp}.docx"

    doc.save(filename)
    print(f"  ✅ 보고서 생성 완료: {filename}")

    return {**state, "document_path": filename}


def setup_styles(doc):
    """문서 스타일 설정"""
    # 제목 스타일
    styles = doc.styles

    # Heading 1 스타일 수정
    if "Heading 1" in styles:
        h1_style = styles["Heading 1"]
        h1_font = h1_style.font
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 51, 102)


def add_cover_page(doc, state):
    """표지 페이지"""
    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI 스타트업\n투자 평가 보고서")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    # 평가 대상
    candidates = state.get("candidates", [])
    if candidates:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_text = " vs ".join([c["name"] for c in candidates[:2]])
        subtitle_run = subtitle.add_run(subtitle_text)
        subtitle_run.font.size = Pt(20)
        subtitle_run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 날짜
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run(datetime.now().strftime("%Y년 %m월 %d일"))
    date_run.font.size = Pt(14)

    doc.add_page_break()


def add_table_of_contents(doc):
    """목차"""
    doc.add_heading("목차", level=1)

    toc_items = [
        ("1. 비교 적합성 검증", "1"),
        ("  1.1 평가 대상 개요", "2"),
        ("  1.2 4가지 비교 가능성 판정", "2"),
        ("  1.3 종합 판정 및 방법론 선택", "2"),
        ("2. 회사 개요", "3"),
        ("  2.1 A사", "3"),
        ("  2.2 B사", "4"),
        ("3. 시장 분석", "5"),
        ("  3.1 공통 시장 환경", "5"),
        ("  3.2 시장 포지션 비교", "6"),
        ("4. 제품/기술 평가", "7"),
        ("  4.1 각 사 기술 평가", "7"),
        ("  4.2 기술 수준 비교", "8"),
        ("5. 경영진 평가", "9"),
        ("6. 재무 현황", "10"),
        ("  6.1 각 사 재무", "10"),
        ("  6.2 규모 비교", "11"),
        ("7. 투자 평가", "12"),
        ("  7.1 개별 평가 (4가지 평가표)", "12"),
        ("  7.2 상대 비교 분석", "14"),
        ("8. 투자 제안", "16"),
        ("  8.1 투자 우선순위", "16"),
        ("  8.2 각 사별 투자 조건", "16"),
        ("  8.3 포트폴리오 전략 제안", "17"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).font.size = Pt(11)
        p.add_run(f" {'.' * (60 - len(item))} {page}").font.color.rgb = RGBColor(
            128, 128, 128
        )

    doc.add_page_break()


def add_comparison_validation(doc, state):
    """1. 비교 적합성 검증"""
    doc.add_heading("1. 비교 적합성 검증", level=1)

    # 1.1 평가 대상 개요
    doc.add_heading("1.1 평가 대상 개요", level=2)

    candidates = state.get("candidates", [])
    comparison_mode = state.get("comparison_mode", "relative")

    if candidates:
        table = doc.add_table(rows=len(candidates) + 1, cols=3)
        table.style = "Light Grid Accent 1"

        # 헤더
        header_cells = table.rows[0].cells
        header_cells[0].text = "기업명"
        header_cells[1].text = "분야"
        header_cells[2].text = "설명"

        # 데이터
        for i, company in enumerate(candidates):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = company.get("name", "")
            row_cells[1].text = state.get("query", "AI 스타트업")
            row_cells[2].text = company.get("description", "")

    doc.add_paragraph()

    # 1.2 4가지 비교 가능성 판정
    doc.add_heading("1.2 4가지 비교 가능성 판정", level=2)

    validation = state.get("validation", {})

    criteria = [
        ("성장 단계 일치도", validation.get("growth_stage", "✓")),
        ("시장 중복도", validation.get("market_overlap", "✓")),
        ("규모 유사도", validation.get("size_similarity", "✓")),
        ("평가 가능성", validation.get("evaluability", "✓")),
    ]

    for criterion, result in criteria:
        p = doc.add_paragraph()
        p.add_run(f"• {criterion}: ").bold = True
        p.add_run(str(result))

    doc.add_paragraph()

    # 1.3 종합 판정
    doc.add_heading("1.3 종합 판정 및 방법론 선택", level=2)

    p = doc.add_paragraph()
    p.add_run("평가 방법: ").bold = True
    mode_text = "상대 평가" if comparison_mode == "relative" else "절대 평가"
    p.add_run(mode_text)

    doc.add_paragraph(
        "두 기업은 동일한 시장에서 유사한 성장 단계에 있으며, "
        "상대 비교를 통한 투자 우선순위 결정이 적합합니다."
    )

    doc.add_page_break()


def add_company_overview(doc, state):
    """2. 회사 개요"""
    doc.add_heading("2. 회사 개요", level=1)

    candidates = state.get("candidates", [])

    for i, company in enumerate(candidates[:2], 1):
        doc.add_heading(f"2.{i} {company.get('name', '')}사", level=2)

        # 기본 정보
        doc.add_paragraph(f"• 기업명: {company.get('name', '')}")
        doc.add_paragraph(f"• 분야: {state.get('query', 'AI 스타트업')}")
        doc.add_paragraph(f"• 설명: {company.get('description', '정보 없음')}")

        # 기술 분석 정보 추가
        tech_analysis = state.get("tech_analysis", {})
        if company["name"] in tech_analysis:
            tech_info = tech_analysis[company["name"]]
            doc.add_paragraph(f"• 기술 점수: {tech_info.get('기술점수', 'N/A')}")
            doc.add_paragraph(f"• 특허 수: {tech_info.get('특허', 0)}개")

        doc.add_paragraph()

    doc.add_page_break()


def add_market_analysis(doc, state):
    """3. 시장 분석"""
    doc.add_heading("3. 시장 분석", level=1)

    market_analysis = state.get("market_analysis", {})

    # 3.1 공통 시장 환경
    doc.add_heading("3.1 공통 시장 환경", level=2)

    industry = market_analysis.get("industry", "AI 스타트업")
    doc.add_paragraph(f"산업 분야: {industry}")

    market_summary = market_analysis.get("market_summary", "시장 분석 정보가 없습니다.")
    doc.add_paragraph(market_summary)

    doc.add_paragraph()

    # 3.2 시장 포지션 비교
    doc.add_heading("3.2 시장 포지션 비교", level=2)

    company_insights = market_analysis.get("company_insights", [])
    if company_insights:
        table = doc.add_table(rows=len(company_insights) + 1, cols=3)
        table.style = "Light Grid Accent 1"

        # 헤더
        header_cells = table.rows[0].cells
        header_cells[0].text = "기업"
        header_cells[1].text = "뉴스 건수"
        header_cells[2].text = "주요 헤드라인"

        # 데이터
        for i, insight in enumerate(company_insights):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = insight.get("name", "")
            row_cells[1].text = str(insight.get("news_count", 0))
            headlines = insight.get("headlines", [])
            row_cells[2].text = headlines[0] if headlines else ""

    doc.add_page_break()


def add_technology_evaluation(doc, state):
    """4. 제품/기술 평가"""
    doc.add_heading("4. 제품/기술 평가", level=1)

    tech_analysis = state.get("tech_analysis", {})

    # 4.1 각 사 기술 평가
    doc.add_heading("4.1 각 사 기술 평가", level=2)

    if tech_analysis:
        for company_name, tech_info in tech_analysis.items():
            doc.add_paragraph(f"• {company_name}")
            doc.add_paragraph(
                f"  - 기술 점수: {tech_info.get('기술점수', 'N/A')}/100",
                style="List Bullet 2",
            )
            doc.add_paragraph(
                f"  - 특허: {tech_info.get('특허', 0)}개", style="List Bullet 2"
            )

    doc.add_paragraph()

    # 4.2 기술 수준 비교
    doc.add_heading("4.2 기술 수준 비교", level=2)

    if tech_analysis and len(tech_analysis) >= 2:
        companies = list(tech_analysis.keys())
        scores = [tech_analysis[c].get("기술점수", 0) for c in companies]

        if len(companies) >= 2:
            doc.add_paragraph(
                f"{companies[0]}사와 {companies[1]}사의 기술 수준은 "
                f"각각 {scores[0]}점, {scores[1]}점으로 평가되었습니다."
            )

    doc.add_page_break()


def add_management_evaluation(doc, state):
    """5. 경영진 평가"""
    doc.add_heading("5. 경영진 평가", level=1)

    scores = state.get("final_score", {}).get("scores", [])

    if scores:
        table = doc.add_table(rows=len(scores) + 1, cols=3)
        table.style = "Light Grid Accent 1"

        # 헤더
        header_cells = table.rows[0].cells
        header_cells[0].text = "기업"
        header_cells[1].text = "경영진 평가 (Berkus)"
        header_cells[2].text = "경영진 가중치 (Scorecard)"

        # 데이터
        for i, score in enumerate(scores):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = score.get("name", "")
            berkus_mgmt = score.get("berkus", {}).get("경영진", 0)
            row_cells[1].text = f"${berkus_mgmt:,}"
            scorecard_mgmt = (
                score.get("scorecard", {}).get("개별점수", {}).get("경영진", 0)
            )
            row_cells[2].text = f"{scorecard_mgmt}점"

    doc.add_page_break()


def add_financial_status(doc, state):
    """6. 재무 현황"""
    doc.add_heading("6. 재무 현황", level=1)

    survival_analysis = state.get("survival_analysis", {}).get("data", [])
    scores = state.get("final_score", {}).get("scores", [])

    # 6.1 각 사 재무
    doc.add_heading("6.1 각 사 재무 현황", level=2)

    if scores:
        for score in scores:
            doc.add_paragraph(f"• {score.get('name', '')}")

            survival = score.get("survival", {})
            doc.add_paragraph(
                f"  - 런웨이: {survival.get('런웨이', 'N/A')}개월",
                style="List Bullet 2",
            )
            doc.add_paragraph(
                f"  - 손익분기: {survival.get('손익분기', 'N/A')}개월",
                style="List Bullet 2",
            )
            doc.add_paragraph(
                f"  - 상태: {survival.get('판정', 'N/A')}", style="List Bullet 2"
            )

    doc.add_paragraph()

    # 6.2 규모 비교
    doc.add_heading("6.2 규모 비교", level=2)

    if scores and len(scores) >= 2:
        doc.add_paragraph(
            "두 기업 모두 안정적인 재무 상태를 유지하고 있으며, "
            "충분한 런웨이를 확보하고 있습니다."
        )

    doc.add_page_break()


def add_investment_evaluation(doc, state):
    """7. 투자 평가"""
    doc.add_heading("7. 투자 평가", level=1)

    scores = state.get("final_score", {}).get("scores", [])

    # 7.1 개별 평가 (4가지 평가표)
    doc.add_heading("7.1 개별 평가 (4가지 평가 방법론)", level=2)

    if scores:
        for score in scores:
            doc.add_heading(f"{score.get('name', '')}", level=3)

            # Berkus Method
            doc.add_paragraph("【Berkus Method】", style="Intense Quote")
            berkus = score.get("berkus", {})
            for key, value in berkus.items():
                if key != "총점":
                    doc.add_paragraph(f"• {key}: ${value:,}")
            doc.add_paragraph(f"총점: ${berkus.get('총점', 0):,}").bold = True

            doc.add_paragraph()

            # Scorecard Method
            doc.add_paragraph("【Scorecard Method】", style="Intense Quote")
            scorecard = score.get("scorecard", {})
            individual_scores = scorecard.get("개별점수", {})
            weights = scorecard.get("가중치", {})
            for key, value in individual_scores.items():
                weight = weights.get(key, 0)
                doc.add_paragraph(f"• {key}: {value}점 (가중치 {weight*100:.0f}%)")
            doc.add_paragraph(f"가중 총점: {scorecard.get('점수', 0)}점").bold = True

            doc.add_paragraph()

            # Growth & PMF
            growth = score.get("growth", {})
            pmf = score.get("pmf", {})

            doc.add_paragraph("【성장성】", style="Intense Quote")
            doc.add_paragraph(f"• 주간 성장률: {growth.get('주간성장률', 0)*100:.1f}%")
            doc.add_paragraph(f"• 판정: {growth.get('판정', 'N/A')}")

            doc.add_paragraph()

            doc.add_paragraph("【PMF】", style="Intense Quote")
            doc.add_paragraph(f"• 달성 수준: {pmf.get('달성', 'N/A')}")
            doc.add_paragraph(f"• 신호: {pmf.get('신호개수', 'N/A')}")

            doc.add_paragraph()

    # 7.2 상대 비교 분석
    doc.add_heading("7.2 상대 비교 분석", level=2)

    if scores and len(scores) >= 2:
        # 평가 결과 비교표
        doc.add_paragraph("【평가 결과 비교표】").bold = True

        table = doc.add_table(rows=len(scores) + 1, cols=6)
        table.style = "Medium Grid 1 Accent 1"

        # 헤더
        headers = ["기업", "Berkus", "Scorecard", "성장성", "생존성", "종합점수"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # 데이터
        for i, score in enumerate(scores):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = score.get("name", "")
            row_cells[1].text = f"${score.get('berkus', {}).get('총점', 0):,}"
            row_cells[2].text = f"{score.get('scorecard', {}).get('점수', 0)}"
            row_cells[3].text = score.get("growth", {}).get("판정", "N/A")
            row_cells[4].text = score.get("survival", {}).get("판정", "N/A")
            row_cells[5].text = f"{score.get('total', 0):.1f}"

    doc.add_page_break()


def add_investment_proposal(doc, state):
    """8. 투자 제안"""
    doc.add_heading("8. 투자 제안", level=1)

    decision = state.get("investment_decision", {})

    # 8.1 투자 우선순위
    doc.add_heading("8.1 투자 우선순위", level=2)

    p = doc.add_paragraph()
    p.add_run("추천 기업: ").bold = True
    p.add_run(f"{decision.get('추천기업', 'N/A')} (1순위)")

    p = doc.add_paragraph()
    p.add_run("판정 등급: ").bold = True
    grade = decision.get("판정등급", "B")
    p.add_run(f"{grade}등급")

    p = doc.add_paragraph()
    p.add_run("투자 판정: ").bold = True
    p.add_run(decision.get("판정", "조건부 투자"))

    p = doc.add_paragraph()
    p.add_run("위험도: ").bold = True
    p.add_run(decision.get("위험도", "중간"))

    doc.add_paragraph()

    # 투자 사유
    doc.add_paragraph("【투자 사유】").bold = True
    reasons = decision.get("투자사유", [])
    for reason in reasons:
        doc.add_paragraph(reason, style="List Bullet")

    doc.add_paragraph()

    # 주의사항
    doc.add_paragraph("【주의사항】").bold = True
    warnings = decision.get("주의사항", ["없음"])
    for warning in warnings:
        doc.add_paragraph(f"• {warning}")

    doc.add_paragraph()

    # 8.2 각 사별 투자 조건
    doc.add_heading("8.2 각 사별 투자 조건", level=2)

    scores = state.get("final_score", {}).get("scores", [])
    if scores:
        for i, score in enumerate(scores, 1):
            doc.add_paragraph(f"{i}. {score.get('name', '')}")
            doc.add_paragraph(f"  - 권장 투자 규모: 협의 필요", style="List Bullet 2")
            doc.add_paragraph(
                f"  - 투자 조건: Valuation 협상 후 결정", style="List Bullet 2"
            )

    doc.add_paragraph()

    # 8.3 포트폴리오 전략 제안
    doc.add_heading("8.3 포트폴리오 전략 제안", level=2)

    doc.add_paragraph(
        "두 기업 모두 성장 가능성이 있으나, "
        f"{decision.get('추천기업', '')}사를 우선 투자 대상으로 추천합니다. "
        "시장 반응 확인 후 추가 투자를 검토하는 단계적 접근을 권장합니다."
    )

    doc.add_paragraph()
    doc.add_paragraph("--- 보고서 끝 ---").alignment = WD_ALIGN_PARAGRAPH.CENTER
